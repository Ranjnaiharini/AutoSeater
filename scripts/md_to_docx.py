from docx import Document
from docx.shared import Pt, Inches
import re
import os

md_path = os.path.join(os.path.dirname(__file__), '..', 'docs', 'PROJECT_REPORT.md')
out_path = os.path.join(os.path.dirname(__file__), '..', 'docs', 'PROJECT_REPORT.docx')

def add_paragraph_with_font(doc, text, font_name='Times New Roman', size=12, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    return p


def convert():
    doc = Document()
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    in_code = False
    code_lines = []

    with open(md_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.rstrip('\n')
            if line.strip().startswith('```'):
                if not in_code:
                    in_code = True
                    code_lines = []
                else:
                    # flush code
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    in_code = False
                continue
            if in_code:
                code_lines.append(line)
                continue

            # Heading
            m = re.match(r'^(#{1,6})\s*(.*)', line)
            if m:
                level = len(m.group(1))
                text = m.group(2).strip()
                if text == '':
                    continue
                # docx heading levels: 0..4 -> use min(level,4)
                hl = min(level,4)
                doc.add_heading(text, level=hl)
                continue

            # Horizontal rule
            if re.match(r'^---+$', line.strip()):
                doc.add_paragraph('')
                continue

            # Image ![alt](path)
            m = re.match(r'^!\[(.*?)\]\((.*?)\)', line.strip())
            if m:
                alt, path = m.group(1), m.group(2)
                # resolve relative path
                img_path = os.path.join(os.path.dirname(os.path.dirname(md_path)), path)
                if os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(6))
                    except Exception:
                        add_paragraph_with_font(doc, f'[Image: {path}]')
                else:
                    add_paragraph_with_font(doc, f'[Missing image: {path}]')
                continue

            # Bullet list
            if re.match(r'^\s*[-*+]\s+(.+)', line):
                item = re.sub(r'^\s*[-*+]\s+', '', line)
                p = doc.add_paragraph(item, style='List Bullet')
                continue

            # Numbered list
            if re.match(r'^\s*\d+\.\s+(.+)', line):
                item = re.sub(r'^\s*\d+\.\s+', '', line)
                p = doc.add_paragraph(item, style='List Number')
                continue

            # Empty line -> paragraph break
            if line.strip() == '':
                doc.add_paragraph('')
                continue

            # Regular paragraph
            add_paragraph_with_font(doc, line)

    doc.save(out_path)
    print('Wrote:', out_path)

if __name__ == '__main__':
    if not os.path.exists(md_path):
        print('Markdown file not found:', md_path)
    else:
        convert()
